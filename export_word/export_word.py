from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def define_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


def export_annotations_to_docx(title, all_annotations, output_path):
    doc = Document()
    define_style(doc)

    doc.add_heading(title, 1)

    for page_num in sorted(all_annotations.keys()):
        doc.add_heading(f'Page {page_num}', level=2)
        doc.add_paragraph()

        for annotation in all_annotations[page_num]:
            for sub in annotation:
                if isinstance(sub, dict): #if dict -> hightlighted annotations
                    resume = sub.get("resume")
                    content = sub.get("content")
                    if resume and content:
                        p = doc.add_paragraph(style='List Bullet')
                        run_resume = p.add_run(resume.strip())
                        run_resume.bold = True
                        run_resume.add_break()
                        p.add_run(content.strip())
                    elif resume:
                        doc.add_paragraph(resume.strip(), style='List Bullet')
                    elif content:
                        doc.add_paragraph(content.strip(), style='List Bullet')
                else: # if text annotations
                    doc.add_paragraph(sub.strip(), style='List Bullet')

        doc.add_paragraph()

    doc.save(output_path)
